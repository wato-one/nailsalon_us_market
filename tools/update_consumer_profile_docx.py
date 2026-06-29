#!/usr/bin/env python3
"""Rebuild Consumer Profile section in the market highlights docx."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = Path(
    "/Users/chinhnguyen/Library/CloudStorage/GoogleDrive-chinhnguyenhue@gmail.com/My Drive"
    "/13.MI.Research/Research & Consultant Companies/Nail Salon Research"
    "/Nail and Spa Industry Consumer Profile & Market Highlights.docx"
)

CA_DEMO = [
    "Population ~39.4M — largest U.S. nail service market; diverse metros (LA, SF, SD, Sacramento).",
    "Core service customers: women ~90%+ (IBISWorld/NAILS); growing male and gender-neutral segment (~8–12%).",
    "Median age ~38; primary demand from ages 25–55 (maintenance) and 65+ (pedicure/spa).",
    "High ethnic diversity — strong nail art and mid-range demand across Hispanic and Asian communities.",
    "Median household income ~USD 100K statewide with coastal premium vs inland price sensitivity.",
    "West region nail product penetration 62% (Statista n=1,092) — proxy for CA adoption.",
]

CA_PSYCHO = [
    "BBC California sets high hygiene expectations (pedicure logs, disinfection) — trust deal-breaker.",
    "Recurring self-care every 2–4 weeks; gel maintenance drives urgency beyond special events.",
    "Coastal premiumization (USD 80–150+ spa) vs inland value and walk-in convenience.",
    "Eco-conscious clients pay more for low-VOC, well-ventilated, 'clean nail' positioning.",
    "~71% won't consider salons under 3★ (BrightLocal) — Google/Yelp filter is mandatory.",
]

SD_DEMO = [
    "San Diego County ~3.3M; City of San Diego ~1.4M (median age 36.6, 54% renters).",
    "East County suburbs (e.g. Santee ~59K) — family-oriented, 70% married-couple households.",
    "Large military population — frequent relocations; hyper-local discovery within 5–15 miles.",
    "County ethnicity: ~41% White, 13% Asian, 35% Hispanic — diverse service preferences.",
    "Median HH income USD 109–113K — supports USD 25–45 manicure and USD 60–80 premium pedicure.",
]

SD_PSYCHO = [
    "Hyper-local discovery — distance and neighborhood reputation beat national brand.",
    "Hygiene and BBC compliance heavily cited in reviews; footspa cleanliness is decisive.",
    "Emotional loyalty to preferred technicians despite many nearby alternatives.",
    "Family weekend visits (East County) require kid-safe, low-odor, clean environments.",
    "Weekend walk-in wait anxiety; ~46–50% of bookings happen outside business hours.",
]

INTRO = (
    "Service-user portrait by market (not nail workforce data). "
    "Demographic = who buys manicure/pedicure/spa nail; "
    "Psychographic = mindset when booking and in-chair."
)


def clear_all_content(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        body.remove(child)


def add_block(doc: Document, blocks: list[tuple[str, str]]) -> None:
    for kind, text in blocks:
        if kind == "h1":
            doc.add_paragraph(text, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(text, style="Heading 2")
        elif kind == "h3":
            doc.add_paragraph(text, style="Heading 3")
        elif kind == "bullet":
            doc.add_paragraph(f"• {text}", style="normal")
        else:
            doc.add_paragraph(text, style="normal")


def main() -> None:
    src = Document(str(DOCX_PATH))
    table_el = deepcopy(src.tables[0]._tbl) if src.tables else None

    highlights = [p.text for p in src.paragraphs[2:5]]
    seg_intro = next(
        p.text
        for p in src.paragraphs
        if p.text.strip().startswith("The consumer base can be divided")
    )
    media = [
        p.text
        for p in src.paragraphs
        if p.text.strip().startswith("Visual Discovery:")
        or p.text.strip().startswith("Reputation & Trust:")
    ]

    consumer_blocks: list[tuple[str, str]] = [
        ("h2", "Consumer Profile"),
        ("normal", INTRO),
        ("h3", "California — Demographic"),
        *(("bullet", t) for t in CA_DEMO),
        ("h3", "California — Psychographic"),
        *(("bullet", t) for t in CA_PSYCHO),
        ("h3", "San Diego — Demographic"),
        *(("bullet", t) for t in SD_DEMO),
        ("h3", "San Diego — Psychographic"),
        *(("bullet", t) for t in SD_PSYCHO),
    ]

    clear_all_content(src)
    src.add_paragraph("Nail and Spa Industry Consumer Profile & Market Highlights", style="Heading 1")
    src.add_paragraph("Market Industry Highlights", style="Heading 2")
    for t in highlights:
        src.add_paragraph(t, style="normal")
    add_block(src, consumer_blocks)
    src.add_paragraph("Consumer Profile Segmentation", style="Heading 2")
    src.add_paragraph(seg_intro, style="normal")
    if table_el is not None:
        src.element.body.append(table_el)
    src.add_paragraph("Media & Discovery Insights", style="Heading 2")
    for t in media:
        src.add_paragraph(t, style="normal")

    src.save(str(DOCX_PATH))
    print(f"Rebuilt: {DOCX_PATH}")


if __name__ == "__main__":
    main()
